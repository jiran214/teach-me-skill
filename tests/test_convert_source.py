#!/usr/bin/env python3
"""测试 convert_source.py 对各种格式的支持。"""

import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent / "scripts"))

from convert_source import (
    convert_to_markdown,
    is_markdown_file,
    parse_file_arg,
)


# ── fixtures ─────────────────────────────────────────────────────


@pytest.fixture
def sample_pdf(tmp_path) -> Path:
    """生成一份简单的测试 PDF。"""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 15, "Test Title", ln=True)
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 12, "Section One", ln=True)
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6, "This is the first section body text.")
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 12, "Section Two", ln=True)
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6, "This is the second section body text.")

    path = tmp_path / "test.pdf"
    pdf.output(str(path))
    return path


@pytest.fixture
def sample_epub(tmp_path) -> Path:
    """生成一份简单的测试 EPUB。"""
    import ebooklib
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier("test-epub-001")
    book.set_title("Test EPUB Book")
    book.set_language("en")
    book.add_author("Test Author")

    ch1 = epub.EpubHtml(title="Chapter 1", file_name="ch1.xhtml", lang="en")
    ch1.content = """
    <html><body>
    <h1>Chapter 1: Getting Started</h1>
    <p>This is the first chapter of the test book.</p>
    <h2>1.1 Background</h2>
    <p>Some background information here.</p>
    <ul><li>Point A</li><li>Point B</li></ul>
    </body></html>
    """

    ch2 = epub.EpubHtml(title="Chapter 2", file_name="ch2.xhtml", lang="en")
    ch2.content = """
    <html><body>
    <h1>Chapter 2: Advanced Topics</h1>
    <p>This is the second chapter.</p>
    <ol><li>First item</li><li>Second item</li></ol>
    </body></html>
    """

    book.add_item(ch1)
    book.add_item(ch2)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", ch1, ch2]

    path = tmp_path / "test.epub"
    epub.write_epub(str(path), book)
    return path


@pytest.fixture
def sample_docx(tmp_path) -> Path:
    """生成一份简单的测试 DOCX。"""
    from docx import Document

    doc = Document()
    doc.add_heading("Test Document Title", level=0)
    doc.add_heading("Chapter 1: Introduction", level=1)
    doc.add_paragraph("This is the introduction paragraph.")
    doc.add_heading("1.1 Overview", level=2)
    doc.add_paragraph("An overview with ")
    run = doc.add_paragraph().add_run("bold text")
    run.bold = True
    doc.add_paragraph("Normal paragraph after bold.")

    doc.add_paragraph("Item one", style="List Bullet")
    doc.add_paragraph("Item two", style="List Bullet")

    doc.add_heading("Chapter 2: Details", level=1)
    doc.add_paragraph("Details go here.")

    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def sample_markdown(tmp_path) -> Path:
    content = """# Sample Markdown

## Section A

This is section A content.

## Section B

This is section B content.
"""
    path = tmp_path / "sample.md"
    path.write_text(content, encoding="utf-8")
    return path


# ── PDF 测试 ─────────────────────────────────────────────────────


class TestPDFConversion:
    def test_basic_structure(self, sample_pdf):
        md = convert_to_markdown(sample_pdf)
        assert "# Test Title" in md
        assert "## Section One" in md
        assert "second section body text" in md

    def test_body_text_present(self, sample_pdf):
        md = convert_to_markdown(sample_pdf)
        assert "first section body text" in md
        assert "second section body text" in md

    def test_returns_string(self, sample_pdf):
        md = convert_to_markdown(sample_pdf)
        assert isinstance(md, str)
        assert len(md) > 0


# ── EPUB 测试（通过 markitdown）──────────────────────────────────


class TestEPUBConversion:
    def test_chapter_headings(self, sample_epub):
        md = convert_to_markdown(sample_epub)
        assert "Chapter 1" in md
        assert "Chapter 2" in md

    def test_subsection_headings(self, sample_epub):
        md = convert_to_markdown(sample_epub)
        assert "1.1 Background" in md

    def test_body_text(self, sample_epub):
        md = convert_to_markdown(sample_epub)
        assert "first chapter of the test book" in md
        assert "second chapter" in md

    def test_list_items(self, sample_epub):
        md = convert_to_markdown(sample_epub)
        assert "Point A" in md
        assert "Point B" in md
        assert "First item" in md


# ── DOCX 测试（通过 markitdown）──────────────────────────────────


class TestDOCXConversion:
    def test_headings(self, sample_docx):
        md = convert_to_markdown(sample_docx)
        assert "Chapter 1" in md
        assert "Chapter 2" in md

    def test_paragraphs(self, sample_docx):
        md = convert_to_markdown(sample_docx)
        assert "introduction paragraph" in md
        assert "Details go here" in md

    def test_list_items(self, sample_docx):
        md = convert_to_markdown(sample_docx)
        assert "Item one" in md
        assert "Item two" in md

    def test_bold_text(self, sample_docx):
        md = convert_to_markdown(sample_docx)
        assert "bold text" in md


# ── Markdown 检测 ────────────────────────────────────────────────


class TestMarkdownDetection:
    def test_md_extension(self):
        assert is_markdown_file(Path("file.md"))

    def test_markdown_extension(self):
        assert is_markdown_file(Path("file.markdown"))

    def test_uppercase(self):
        assert is_markdown_file(Path("FILE.MD"))

    def test_non_markdown(self):
        assert not is_markdown_file(Path("file.pdf"))
        assert not is_markdown_file(Path("file.epub"))
        assert not is_markdown_file(Path("file.docx"))


# ── parse_file_arg 测试 ──────────────────────────────────────────


class TestParseFileArg:
    def test_simple_path(self, tmp_path):
        f = tmp_path / "test.pdf"
        f.touch()
        path, title = parse_file_arg(str(f))
        assert path == f
        assert title is None

    def test_path_with_title(self, tmp_path):
        f = tmp_path / "test.pdf"
        f.touch()
        path, title = parse_file_arg(f"{f}:My Title")
        assert path == f
        assert title == "My Title"

    def test_path_without_colon(self):
        path, title = parse_file_arg("some_file.pdf")
        assert path == Path("some_file.pdf")
        assert title is None

    def test_nonexistent_file_with_colon(self):
        path, title = parse_file_arg("nonexistent.pdf:some title")
        assert path == Path("nonexistent.pdf:some title")
        assert title is None

    def test_empty_title(self, tmp_path):
        f = tmp_path / "test.pdf"
        f.touch()
        path, title = parse_file_arg(f"{f}:")
        assert path == f
        assert title is None


# ── 端到端测试 ────────────────────────────────────────────────────


class TestEndToEnd:
    def test_save_pdf_with_title(self, sample_pdf, tmp_path):
        sources_dir = tmp_path / "sources"
        sources_dir.mkdir()
        md = convert_to_markdown(sample_pdf)
        from convert_source import save_markdown

        out = save_markdown(md, sample_pdf, sources_dir, title="Custom Title")
        content = out.read_text(encoding="utf-8")
        assert content.startswith("# Custom Title")
        assert "Test Title" in content

    def test_save_epub_with_title(self, sample_epub, tmp_path):
        sources_dir = tmp_path / "sources"
        sources_dir.mkdir()
        md = convert_to_markdown(sample_epub)
        from convert_source import save_markdown

        out = save_markdown(md, sample_epub, sources_dir, title="EPUB Title")
        content = out.read_text(encoding="utf-8")
        assert content.startswith("# EPUB Title")
        assert "Chapter 1" in content

    def test_save_docx_with_title(self, sample_docx, tmp_path):
        sources_dir = tmp_path / "sources"
        sources_dir.mkdir()
        md = convert_to_markdown(sample_docx)
        from convert_source import save_markdown

        out = save_markdown(md, sample_docx, sources_dir, title="DOCX Title")
        content = out.read_text(encoding="utf-8")
        assert content.startswith("# DOCX Title")
        assert "Chapter 1" in content
